from docx import Document
import markdown2
from PyPDF2 import PdfReader
import os
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup
import tempfile
from pdf2image import convert_from_path
import pytesseract
import logging
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

# 配置Tesseract路径
tesseract_path = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
if os.path.exists(tesseract_path):
    pytesseract.pytesseract.tesseract_cmd = tesseract_path

logger = logging.getLogger(__name__)

def process_file(file_path: str, file_extension: str) -> str:
    """
    根据文件类型处理文件内容并返回文本
    """
    if file_extension == 'html':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    elif file_extension == 'docx':
        doc = Document(file_path)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    
    elif file_extension == 'md':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    elif file_extension == 'pdf':
        reader = PdfReader(file_path)
        if len(reader.pages) == 0:
            raise ValueError(f"PDF文件 '{os.path.basename(file_path)}' 没有任何页面")
            
        text = ''
        for page_num, page in enumerate(reader.pages, 1):
            try:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + '\n'
            except Exception as e:
                raise ValueError(f"无法从PDF文件 '{os.path.basename(file_path)}' 的第 {page_num} 页提取文本: {str(e)}")
                
        if not text.strip():
            logger.info("PDF文件没有可直接提取的文本内容，尝试使用OCR识别...")
            text = extract_text_from_pdf_with_ocr(file_path)
            
        return text
    
    elif file_extension == 'epub':
        book = epub.read_epub(file_path)
        text = ''
        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                soup = BeautifulSoup(item.get_content(), 'html.parser')
                text += soup.get_text() + '\n'
        return text
    
    raise ValueError(f"Unsupported file format: {file_extension}")

def extract_text_from_pdf_with_ocr(pdf_path: str) -> str:
    """
    使用OCR从PDF文件中提取文本
    """
    try:
        # 检查Tesseract是否正确配置
        if not os.path.exists(tesseract_path):
            raise ValueError("Tesseract-OCR未安装或路径不正确。请安装Tesseract-OCR并确保安装在正确的位置。")

        logger.info(f"开始使用OCR处理PDF文件: {pdf_path}")
        # 将PDF转换为图片
        images = convert_from_path(pdf_path)
        
        # 使用OCR处理每一页
        text = ""
        for i, image in enumerate(images, 1):
            logger.info(f"正在处理第 {i}/{len(images)} 页")
            # 对图片进行OCR处理
            page_text = pytesseract.image_to_string(image, lang='eng')
            if page_text.strip():
                text += page_text + "\n\n"
                
        if not text.strip():
            raise ValueError("OCR未能识别出任何文本内容")
            
        return text
        
    except Exception as e:
        raise ValueError(f"OCR处理失败: {str(e)}")

def convert_pdf_to_markdown(pdf_path: str) -> str:
    """
    将PDF文件转换为Markdown格式，如果无法直接提取文本则使用OCR
    """
    try:
        logger.info(f"开始读取PDF文件: {pdf_path}")
        reader = PdfReader(pdf_path)
        markdown_content = []
        need_ocr = True
        
        # 首先尝试直接提取文本
        logger.info("尝试直接提取PDF文本")
        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text.strip():
                need_ocr = False
                break
        
        # 如果无法直接提取文本，使用OCR
        if need_ocr:
            logger.info("直接提取文本失败，尝试使用OCR")
            if not os.path.exists(tesseract_path):
                raise ValueError("Tesseract-OCR未安装或路径不正确。请安装Tesseract-OCR并确保安装在正确的位置。")
            
            # 将PDF转换为图片
            images = convert_from_path(pdf_path)
            logger.info(f"PDF已转换为{len(images)}页图片")
            
            # 使用OCR处理每一页
            for i, image in enumerate(images, 1):
                logger.info(f"正在OCR处理第 {i}/{len(images)} 页")
                
                # 对图片进行OCR处理
                page_text = pytesseract.image_to_string(image, lang='eng+chi')
                
                if page_text.strip():
                    # 添加页码标记
                    markdown_content.append(f"\n## Page {i}\n")
                    
                    # 处理OCR文本
                    paragraphs = page_text.split('\n\n')
                    for paragraph in paragraphs:
                        if paragraph.strip():
                            # 处理每一行
                            lines = paragraph.split('\n')
                            for line in lines:
                                line = line.strip()
                                if not line:
                                    continue
                                    
                                # 检测标题（全大写或以#开头）
                                if line.isupper() or line.startswith('#'):
                                    markdown_content.append(f"\n### {line}\n")
                                # 检测列表项
                                elif line.startswith(('•', '-', '*', '○', '>')):
                                    markdown_content.append(f"- {line[1:].strip()}")
                                # 普通段落
                                else:
                                    markdown_content.append(line)
                            
                            markdown_content.append("\n")  # 段落之间添加空行
                else:
                    logger.warning(f"第 {i} 页OCR未识别出文本")
        
        # 直接提取文本
        else:
            logger.info("使用直接提取文本的方式")
            for page_num, page in enumerate(reader.pages, 1):
                logger.info(f"正在处理第 {page_num} 页")
                text = page.extract_text()
                
                if not text.strip():
                    logger.warning(f"第 {page_num} 页没有文本内容")
                    continue
                
                # 添加页码标记
                markdown_content.append(f"\n## Page {page_num}\n")
                
                # 处理段落
                paragraphs = text.split('\n\n')
                for paragraph in paragraphs:
                    if not paragraph.strip():
                        continue
                    
                    lines = paragraph.split('\n')
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                        
                        # 检测标题
                        if len(line) < 100 and line.isupper():
                            markdown_content.append(f"\n### {line}\n")
                        # 检测列表项
                        elif line.startswith(('•', '-', '*', '○')):
                            markdown_content.append(f"- {line[1:].strip()}")
                        # 普通段落
                        else:
                            markdown_content.append(line)
                    
                    markdown_content.append("\n")  # 段落之间添加空行
        
        final_content = "\n".join(markdown_content)
        logger.info(f"PDF转换Markdown完成，内容长度: {len(final_content)}")
        
        if not final_content.strip():
            raise ValueError("未能提取到任何文本内容")
            
        return final_content
        
    except Exception as e:
        logger.error(f"PDF转Markdown失败: {str(e)}", exc_info=True)
        raise ValueError(f"无法将PDF转换为Markdown: {str(e)}")

def save_translated_file(translated_text: str, output_path: str, file_extension: str, original_file_path: str = None):
    """
    将翻译后的文本保存为对应格式的文件
    """
    try:
        logger.info(f"开始保存文件: {output_path}")
        logger.info(f"文件格式: {file_extension}")
        
        # 确保输出目录存在
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        if file_extension in ['html', 'md', 'txt']:
            logger.info(f"保存为文本文件: {output_path}")
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(translated_text)
            logger.info(f"文本文件保存成功: {output_path}")
            return output_path
        
        elif file_extension == 'docx':
            doc = Document()
            for paragraph in translated_text.split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph)
            doc.save(output_path)
            return output_path
        
        elif file_extension == 'pdf':
            # 创建PDF文档
            doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            # 注册中文字体
            try:
                # 尝试使用系统中文字体
                font_paths = [
                    "C:/Windows/Fonts/simhei.ttf",  # 黑体
                    "C:/Windows/Fonts/simsun.ttc",  # 宋体
                    "C:/Windows/Fonts/msyh.ttc"     # 微软雅黑
                ]
                
                font_found = False
                for font_path in font_paths:
                    if os.path.exists(font_path):
                        pdfmetrics.registerFont(TTFont('chinese', font_path))
                        font_found = True
                        break
                        
                if not font_found:
                    logger.warning("未找到中文字体，将使用默认字体")
            except Exception as e:
                logger.error(f"注册中文字体失败: {str(e)}")
            
            # 创建支持中文的样式
            styles = getSampleStyleSheet()
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontName='chinese' if font_found else 'Helvetica',
                fontSize=12,
                leading=14,
                wordWrap='CJK'  # 支持中文换行
            )
            
            # 将文本分段
            paragraphs = []
            for text in translated_text.split('\n'):
                if text.strip():
                    # 处理特殊字符
                    text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    paragraphs.append(Paragraph(text, normal_style))
                    paragraphs.append(Spacer(1, 12))  # 添加段落间距
            
            # 构建PDF
            try:
                doc.build(paragraphs)
                return output_path
            except Exception as e:
                logger.error(f"PDF生成失败: {str(e)}")
                # 如果PDF生成失败，回退到文本文件
                txt_path = output_path.rsplit('.', 1)[0] + '.txt'
                with open(txt_path, 'w', encoding='utf-8') as f:
                    f.write(translated_text)
                return txt_path
                
        elif file_extension == 'epub':
            # 创建新的EPUB
            book = epub.EpubBook()
            
            # 如果有原始文件，复制元数据
            if original_file_path:
                try:
                    original_book = epub.read_epub(original_file_path)
                    book.metadata = original_book.metadata
                    book.spine = original_book.spine
                except Exception as e:
                    logger.warning(f"复制原EPUB元数据失败: {str(e)}")
            
            # 创建章节
            c1 = epub.EpubHtml(title='Content',
                            file_name='content.xhtml',
                            content=f'<html><body>{translated_text}</body></html>')
            
            # 添加章节
            book.add_item(c1)
            
            # 创建spine
            book.spine = ['nav', c1]
            
            # 添加默认CSS
            style = 'BODY {color: white;}'
            nav_css = epub.EpubItem(uid="style_nav",
                                file_name="style/nav.css",
                                media_type="text/css",
                                content=style)
            book.add_item(nav_css)
            
            # 创建导航
            book.toc = [epub.Link('content.xhtml', '内容', 'content')]
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            
            # 写入文件
            epub.write_epub(output_path, book, {})
            return output_path
            
        else:
            raise ValueError(f"不支持的文件格式: {file_extension}")
            
    except Exception as e:
        logger.error(f"保存翻译文件失败: {str(e)}", exc_info=True)
        raise ValueError(f"保存翻译文件失败: {str(e)}")
